"""文档解析 — 使用 PyPDF2 + python-docx 逐页/逐段提取文本，支持部分容错

对齐 ARCHITECTURE.md §4.7:
- 单页/单段失败跳过并记录 warning
- < 20% 失败 → 继续（记录 warning）
- 20%~50% 失败 → partial_failed
- > 50% 失败 → failed
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    """单页解析结果"""
    page_number: int
    content: str
    success: bool = True
    error: str | None = None


@dataclass
class ParseResult:
    """文档解析聚合结果"""
    pages: list[ParsedPage] = field(default_factory=list)
    total_pages: int = 0
    failed_pages: int = 0

    @property
    def failure_rate(self) -> float:
        if self.total_pages == 0:
            return 1.0  # 空文档视为全部失败
        return self.failed_pages / self.total_pages

    @property
    def full_text(self) -> str:
        """拼接所有成功页面的文本"""
        return "\n\n".join(p.content for p in self.pages if p.success)

    @property
    def warnings(self) -> list[str]:
        """收集所有失败页面的警告信息"""
        return [
            f"第{p.page_number}页: {p.error}"
            for p in self.pages
            if not p.success and p.error
        ]


def parse_document(file_path: str, file_type: str | None = None) -> ParseResult:
    """解析文档主入口，根据文件类型分发到对应解析器。

    Args:
        file_path: 文档文件绝对路径
        file_type: 文件类型（pdf/docx/md/txt），为 None 时从扩展名推断

    Returns:
        ParseResult: 包含逐页/逐段解析结果和容错统计
    """
    path = Path(file_path)
    if not path.exists():
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=f"文件不存在: {file_path}")],
            total_pages=1,
            failed_pages=1,
        )

    if file_type is None:
        file_type = path.suffix.lower().lstrip(".")

    try:
        if file_type == "pdf":
            return _parse_pdf(file_path)
        elif file_type == "docx":
            return _parse_docx(file_path)
        elif file_type in ("md", "txt"):
            return _parse_text(file_path)
        else:
            return ParseResult(
                pages=[ParsedPage(page_number=1, content="", success=False, error=f"不支持的文件类型: {file_type}")],
                total_pages=1,
                failed_pages=1,
            )
    except Exception as e:
        logger.exception(f"文档解析异常: {file_path}")
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )


def _parse_pdf(file_path: str) -> ParseResult:
    """使用 PyPDF2 逐页解析 PDF，单页失败跳过并记录"""
    try:
        reader = PdfReader(file_path)
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    pages: list[ParsedPage] = []
    failed = 0

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                pages.append(ParsedPage(page_number=i + 1, content=text.strip()))
            else:
                pages.append(ParsedPage(
                    page_number=i + 1, content="",
                    success=False, error="页面无文本或文本为空"
                ))
                failed += 1
        except Exception as e:
            pages.append(ParsedPage(
                page_number=i + 1, content="",
                success=False, error=f"页面解析异常: {e}"
            ))
            failed += 1

    total = len(reader.pages)
    return ParseResult(pages=pages, total_pages=total, failed_pages=failed)


def _parse_docx(file_path: str) -> ParseResult:
    """使用 python-docx 解析 DOCX，逐段提取并容错（对齐 PDF 逐页容错粒度）"""
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    if not doc.paragraphs:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文档无段落内容")],
            total_pages=1,
            failed_pages=1,
        )

    pages: list[ParsedPage] = []
    failed = 0

    for i, p in enumerate(doc.paragraphs):
        try:
            text = p.text
            if text and text.strip():
                pages.append(ParsedPage(page_number=i + 1, content=text.strip()))
        except Exception as e:
            logger.warning(f"DOCX 第{i+1}段解析失败: {e}")
            pages.append(ParsedPage(
                page_number=i + 1, content="",
                success=False, error=f"段落解析异常: {e}"
            ))
            failed += 1

    total = len(doc.paragraphs)

    if not pages:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文档无有效文本内容")],
            total_pages=total,
            failed_pages=total,
        )

    return ParseResult(pages=pages, total_pages=total, failed_pages=failed)


def _parse_text(file_path: str) -> ParseResult:
    """解析纯文本文件（md/txt），统一 UTF-8 读取"""
    try:
        content = Path(file_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = Path(file_path).read_text(encoding="gbk")
        except Exception as e:
            return ParseResult(
                pages=[ParsedPage(page_number=1, content="", success=False, error=f"编码错误: {e}")],
                total_pages=1,
                failed_pages=1,
            )
    except Exception as e:
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error=str(e))],
            total_pages=1,
            failed_pages=1,
        )

    if not content.strip():
        return ParseResult(
            pages=[ParsedPage(page_number=1, content="", success=False, error="文件内容为空")],
            total_pages=1,
            failed_pages=1,
        )

    page = ParsedPage(page_number=1, content=content.strip())
    return ParseResult(pages=[page], total_pages=1, failed_pages=0)
